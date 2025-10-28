import io
import json
import xml.etree.ElementTree as ET
import os
import logging
from pathlib import Path
from datetime import datetime, date
from typing import List, Dict, Any, Optional, BinaryIO
import pandas as pd
import matplotlib
matplotlib.use('Agg') # Usar backend no interactivo
import matplotlib.pyplot as plt

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.chart import BarChart, Reference

from app.schemas import ReporteDiario, ReporteTurno, EventoResponse, TurnoEnum
from app.utils import obtener_nombre_turno, obtener_horario_turno_str

logger = logging.getLogger(__name__)

class ReportGenerator:
    """
    Generador de reportes en múltiples formatos.
    """
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        base_dir = Path(__file__).resolve().parent.parent.parent
        self.logo_path = base_dir / "public" / "logo.png"
        logger.info(f"Ruta del logo construida: {self.logo_path}")
        logger.info(f"¿Existe el logo en la ruta? {'Sí' if self.logo_path.exists() else 'No'}")

        primary_color = colors.HexColor('#6c63ff')

        # Safely add styles to avoid redefinition error
        if 'ReportTitle' not in self.styles:
            self.styles.add(ParagraphStyle(name='ReportTitle', 
                                           parent=self.styles['h1'], 
                                           fontSize=24, 
                                           leading=28, 
                                           alignment=1, 
                                           spaceAfter=20, 
                                           textColor=primary_color))
        if 'SectionHeading' not in self.styles:
            self.styles.add(ParagraphStyle(name='SectionHeading', 
                                           parent=self.styles['h2'], 
                                           fontSize=16, 
                                           leading=18, 
                                           spaceBefore=15, 
                                           spaceAfter=10, 
                                           textColor=primary_color))
        if 'SubSectionHeading' not in self.styles:
            self.styles.add(ParagraphStyle(name='SubSectionHeading', 
                                           parent=self.styles['h3'], 
                                           fontSize=14, 
                                           leading=16, 
                                           spaceBefore=10, 
                                           spaceAfter=5, 
                                           textColor=primary_color))
        if 'Footer' not in self.styles:
            self.styles.add(ParagraphStyle(name='Footer', 
                                           parent=self.styles['Normal'], 
                                           fontSize=9, 
                                           alignment=1, 
                                           textColor=colors.grey, 
                                           spaceBefore=30))

    def _add_logo_to_pdf(self, story):
        if self.logo_path.exists():
            logo = Image(self.logo_path, width=150, height=60)
            logo.hAlign = 'LEFT'
            story.append(logo)
            story.append(Spacer(1, 20))

    def _create_bar_chart(self, data: ReporteDiario) -> io.BytesIO:
        """Crea un gráfico de barras con Matplotlib y lo devuelve como un buffer de BytesIO."""
        labels = [f"Turno {t.turno.value}" for t in data.turnos]
        empleados = [t.total_empleados for t in data.turnos]
        faltas = [t.total_faltas for t in data.turnos]

        x = pd.Series(range(len(labels)))
        width = 0.35
        
        fig, ax = plt.subplots(figsize=(8, 4))
        rects1 = ax.bar(x - width/2, empleados, width, label='Empleados', color='#4299e1')
        rects2 = ax.bar(x + width/2, faltas, width, label='Faltas', color='#f56565')

        ax.set_ylabel('Cantidad')
        ax.set_title('Empleados y Faltas por Turno')
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.legend()

        ax.bar_label(rects1, padding=3)
        ax.bar_label(rects2, padding=3)

        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150)
        buf.seek(0)
        plt.close(fig)
        return buf

    # ==================== GENERACIÓN PDF ====================
    
    def generate_pdf_report(self, data: ReporteDiario, include_charts: bool = True) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        primary_color = colors.HexColor('#6c63ff')

        # --- Header (Logo and Title) ---
        if self.logo_path.exists():
            logo = Image(self.logo_path, width=150, height=60)
            logo.hAlign = 'LEFT'
            story.append(logo)
            story.append(Spacer(1, 10))

        story.append(Paragraph(f"Reporte de Detección de Empleados", self.styles['ReportTitle']))
        story.append(Paragraph(f"<font color=\"#000000\">{data.fecha}</font>", self.styles['SectionHeading']))
        story.append(Spacer(1, 20))
        
        # --- Resumen General ---
        story.append(Paragraph("Resumen General", self.styles['SectionHeading']))
        summary_data = [
            ['Descripción', 'Valor'],
            ['Total Empleados Detectados', str(data.total_empleados)],
            ['Total Faltas Detectadas', str(data.total_faltas)],
            ['Tasa de Faltas', f"{(data.total_faltas/data.total_empleados*100):.1f}%" if data.total_empleados > 0 else "0%"],
            ['Confianza Promedio', f"{(data.promedio_confianza_general or 0):.1f}%" ]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary_color),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 30))

        # --- Gráfico de Barras ---
        if include_charts and any(t.total_empleados > 0 or t.total_faltas > 0 for t in data.turnos):
            story.append(Paragraph("Gráfico de Empleados y Faltas por Turno", self.styles['SubSectionHeading']))
            chart_buffer = self._create_bar_chart(data)
            story.append(Image(chart_buffer, width=6*inch, height=3*inch))
            story.append(Spacer(1, 20))
        
        # --- Reporte por Turno ---
        story.append(Paragraph("Detalle por Turno", self.styles['SectionHeading']))
        for turno_report in data.turnos:
            story.append(Paragraph(f"Turno {turno_report.turno.value}: {turno_report.nombre_turno}", self.styles['SubSectionHeading']))
            story.append(Paragraph(f"Horario: {turno_report.horario}", self.styles['Normal']))
            story.append(Spacer(1, 10))
            
            if turno_report.eventos:
                eventos_data = [['Hora', 'Empleados', 'Faltas', 'Confianza', 'Descripción']]
                for evento in turno_report.eventos[:10]:
                    eventos_data.append([
                        evento.fecha_hora.strftime('%H:%M:%S'),
                        str(evento.empleados),
                        str(evento.faltas),
                        f"{(evento.confianza or 0):.1f}%",
                        evento.descripcion[:50] + "..." if evento.descripcion and len(evento.descripcion) > 50 else evento.descripcion or ""
                    ])
                eventos_table = Table(eventos_data, colWidths=[1*inch, 0.8*inch, 0.8*inch, 0.8*inch, 3*inch])
                eventos_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ]))
                story.append(eventos_table)
            else:
                story.append(Paragraph("No hay eventos registrados.", self.styles['Normal']))
            story.append(Spacer(1, 20))
        
        # --- Footer ---
        story.append(Paragraph("&copy; Sisa sistemas automatizados 2025", self.styles['Footer']))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    # ==================== GENERACIÓN WORD ====================
    
    def generate_word_report(self, data: ReporteDiario, include_charts: bool = True) -> bytes:
        doc = Document()
        
        # --- Header (Logo and Title) ---
        if self.logo_path.exists():
            doc.add_picture(str(self.logo_path), width=Inches(1.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        title = doc.add_heading(f'Reporte de Detección de Empleados', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'{data.fecha}', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
        

        # --- Resumen General ---
        doc.add_heading('Resumen General', level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        
        # Header row
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Descripción'
        hdr_cells[1].text = 'Valor'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        summary_data = [
            ('Total Empleados Detectados', str(data.total_empleados)),
            ('Total Faltas Detectadas', str(data.total_faltas)),
            ('Tasa de Faltas', f"{(data.total_faltas/data.total_empleados*100):.1f}%" if data.total_empleados > 0 else "0%"),
            ('Confianza Promedio', f"{(data.promedio_confianza_general or 0):.1f}%" )
        ]
        for metric, value in summary_data:
            row_cells = summary_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value

        # --- Gráfico de Barras ---
        if include_charts and any(t.total_empleados > 0 or t.total_faltas > 0 for t in data.turnos):
            doc.add_heading('Gráfico de Empleados y Faltas por Turno', level=2)
            chart_buffer = self._create_bar_chart(data)
            doc.add_picture(chart_buffer, width=Inches(6))

        # --- Detalle por Turno ---
        doc.add_heading('Detalle por Turno', level=1)
        for turno_report in data.turnos:
            doc.add_heading(f'Turno {turno_report.turno.value}: {turno_report.nombre_turno}', level=2)
            doc.add_paragraph(f'Horario: {turno_report.horario}')
            if turno_report.eventos:
                eventos_table = doc.add_table(rows=1, cols=5)
                eventos_table.style = 'Table Grid'
                headers = ['Hora', 'Empleados', 'Faltas', 'Confianza', 'Descripción']
                for i, header in enumerate(headers):
                    cell = eventos_table.cell(0, i)
                    cell.text = header
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for evento in turno_report.eventos[:10]:
                    row = eventos_table.add_row().cells
                    row[0].text = evento.fecha_hora.strftime('%H:%M:%S')
                    row[1].text = str(evento.empleados)
                    row[2].text = str(evento.faltas)
                    row[3].text = f"{(evento.confianza or 0):.1f}%"
                    row[4].text = evento.descripcion[:100] if evento.descripcion else ""
            else:
                doc.add_paragraph("No hay eventos registrados.")
        
        # --- Footer ---
        footer = doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "© Sisa sistemas automatizados 2025"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    # ==================== GENERACIÓN EXCEL ====================
    
    def generate_excel_report(self, data: ReporteDiario, include_charts: bool = True) -> bytes:
        wb = Workbook()
        ws_summary = wb.active
        ws_summary.title = "Resumen"

        # --- Estilos Definidos ---
        header_font = Font(size=18, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4A90E2", end_color="4A90E2", fill_type="solid")
        subheader_font = Font(size=14, bold=True)
        subheader_fill = PatternFill(start_color="F0F0F0", end_color="F0F0F0", fill_type="solid")
        total_font = Font(bold=True)
        center_alignment = Alignment(horizontal='center', vertical='center')
        
        # --- Logo y Título ---
        if self.logo_path.exists():
            img = OpenpyxlImage(self.logo_path)
            img.width = 150
            img.height = 60
            ws_summary.add_image(img, 'A1')

        ws_summary.merge_cells('C1:H2')
        title_cell = ws_summary['C1']
        title_cell.value = f'Reporte de Detección - {data.fecha}'
        title_cell.font = header_font
        title_cell.fill = header_fill
        title_cell.alignment = center_alignment

        # --- Resumen General ---
        ws_summary.merge_cells('A4:B4')
        summary_title_cell = ws_summary['A4']
        summary_title_cell.value = 'Resumen General'
        summary_title_cell.font = subheader_font
        summary_title_cell.fill = subheader_fill
        summary_title_cell.alignment = Alignment(horizontal='center')

        summary_data = [
            ('Total Empleados', data.total_empleados),
            ('Total Faltas', data.total_faltas),
            ('Tasa de Faltas', f"{(data.total_faltas/data.total_empleados*100):.1f}%" if data.total_empleados > 0 else "0%"),
            ('Confianza Promedio', f"{(data.promedio_confianza_general or 0):.1f}%")
        ]
        for i, (metric, value) in enumerate(summary_data, start=5):
            ws_summary[f'A{i}'].value = metric
            ws_summary[f'B{i}'].value = value
            ws_summary[f'A{i}'].font = total_font

        # --- Gráfico de Barras ---
        if include_charts and any(t.total_empleados > 0 or t.total_faltas > 0 for t in data.turnos):
            chart_buffer = self._create_bar_chart(data)
            img = OpenpyxlImage(chart_buffer)
            img.anchor = 'D4'
            ws_summary.add_image(img)

        # --- Hoja de Eventos Detallados ---
        ws_events = wb.create_sheet("Eventos Detallados")
        headers = ['Fecha/Hora', 'Turno', 'Empleados', 'Faltas', 'Confianza', 'Descripción']
        for col_idx, header in enumerate(headers, start=1):
            cell = ws_events.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4A90E2", end_color="4A90E2", fill_type="solid")
            cell.alignment = center_alignment

        row_idx = 2
        for turno_report in data.turnos:
            for evento in turno_report.eventos:
                ws_events.cell(row=row_idx, column=1, value=evento.fecha_hora.strftime('%Y-%m-%d %H:%M:%S'))
                ws_events.cell(row=row_idx, column=2, value=f"Turno {evento.turno}")
                ws_events.cell(row=row_idx, column=3, value=evento.empleados)
                ws_events.cell(row=row_idx, column=4, value=evento.faltas)
                ws_events.cell(row=row_idx, column=5, value=f"{(evento.confianza or 0):.1f}%")
                ws_events.cell(row=row_idx, column=6, value=evento.descripcion or "")
                
                # Resaltar faltas
                if evento.faltas > 0:
                    ws_events.cell(row=row_idx, column=4).fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
                    ws_events.cell(row=row_idx, column=4).font = Font(color="9C0006")

                row_idx += 1

        # --- Ajuste de Columnas ---
        for ws in [ws_summary, ws_events]:
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 4, 60)
                ws.column_dimensions[column_letter].width = adjusted_width

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_xml_report(self, data: ReporteDiario) -> str:
        """
        Genera un reporte en formato XML.
        """
        root = ET.Element("ReporteDeteccionEmpleados")
        root.set("fecha", data.fecha)
        root.set("generado", datetime.now().isoformat())
        
        resumen = ET.SubElement(root, "ResumenGeneral")
        ET.SubElement(resumen, "TotalEmpleados").text = str(data.total_empleados)
        ET.SubElement(resumen, "TotalFaltas").text = str(data.total_faltas)
        ET.SubElement(resumen, "TasaFaltas").text = f"{(data.total_faltas/data.total_empleados*100):.1f}" if data.total_empleados > 0 else "0"
        ET.SubElement(resumen, "PromedioConfianza").text = f"{data.promedio_confianza_general:.1f}" if data.promedio_confianza_general else "0"
        
        turnos_elem = ET.SubElement(root, "Turnos")
        for turno_report in data.turnos:
            turno_elem = ET.SubElement(turnos_elem, "Turno")
            turno_elem.set("numero", str(turno_report.turno.value))
            turno_elem.set("nombre", turno_report.nombre_turno)
            turno_elem.set("horario", turno_report.horario)
            ET.SubElement(turno_elem, "TotalEmpleados").text = str(turno_report.total_empleados)
            ET.SubElement(turno_elem, "TotalFaltas").text = str(turno_report.total_faltas)
            ET.SubElement(turno_elem, "PromedioConfianza").text = f"{turno_report.promedio_confianza:.1f}" if turno_report.promedio_confianza else "0"
            
            eventos_elem = ET.SubElement(turno_elem, "Eventos")
            for evento in turno_report.eventos:
                evento_elem = ET.SubElement(eventos_elem, "Evento")
                evento_elem.set("id", str(evento.id_evento))
                ET.SubElement(evento_elem, "FechaHora").text = evento.fecha_hora.isoformat()
                ET.SubElement(evento_elem, "Empleados").text = str(evento.empleados)
                ET.SubElement(evento_elem, "Faltas").text = str(evento.faltas)
                ET.SubElement(evento_elem, "Confianza").text = str(evento.confianza) if evento.confianza else "0"
                ET.SubElement(evento_elem, "Descripcion").text = evento.descripcion or ""
        
        return ET.tostring(root, encoding='unicode', xml_declaration=True)
    
    def generate_json_report(self, data: ReporteDiario) -> str:
        """
        Genera un reporte en formato JSON.
        """
        report_dict = {
            "fecha": data.fecha,
            "generado": datetime.now().isoformat(),
            "resumen_general": {
                "total_empleados": data.total_empleados,
                "total_faltas": data.total_faltas,
                "tasa_faltas": round((data.total_faltas/data.total_empleados*100), 2) if data.total_empleados > 0 else 0,
                "promedio_confianza": round(data.promedio_confianza_general, 2) if data.promedio_confianza_general else None
            },
            "turnos": []
        }
        
        for turno_report in data.turnos:
            turno_dict = {
                "numero": turno_report.turno.value,
                "nombre": turno_report.nombre_turno,
                "horario": turno_report.horario,
                "total_empleados": turno_report.total_empleados,
                "total_faltas": turno_report.total_faltas,
                "promedio_confianza": round(turno_report.promedio_confianza, 2) if turno_report.promedio_confianza else None,
                "fecha_inicio": turno_report.fecha_inicio.isoformat(),
                "fecha_fin": turno_report.fecha_fin.isoformat(),
                "eventos": []
            }
            
            for evento in turno_report.eventos:
                evento_dict = {
                    "id_evento": evento.id_evento,
                    "fecha_hora": evento.fecha_hora.isoformat(),
                    "empleados": evento.empleados,
                    "faltas": evento.faltas,
                    "confianza": evento.confianza,
                    "descripcion": evento.descripcion
                }
                turno_dict["eventos"].append(evento_dict)
            
            report_dict["turnos"].append(turno_dict)
        
        return json.dumps(report_dict, indent=2, ensure_ascii=False)

report_generator = ReportGenerator()